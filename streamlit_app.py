import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import traceback
import sqlite3 # DATABASE LIBRARY

class WordCardGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Al-Ghozali Word Generator (Database Integrated)")
        self.root.geometry("1000x750")
        self.root.configure(bg="#f0fdf4") # Nuansa Hijau Soft

        # --- DATA DEFAULT ---
        self.DATA_SEKOLAH = {
            "SMA ISLAM AL-GHOZALI": "Antoni Firdaus, SHI, M.Pd.",
            "SMP ISLAM AL-GHOZALI": "Iswahyudin, SE"
        }
        self.jadwal_data = [] 
        
        # Inisialisasi Database
        self.init_db()
        
        # Setup UI
        self.setup_ui()
        
        # Load Data Tersimpan setelah UI siap
        self.load_saved_config()

    # ==========================================
    # 1. DATABASE MANAGER (SQLITE)
    # ==========================================
    def init_db(self):
        """Membuat file database.db dan tabel konfigurasi jika belum ada"""
        try:
            self.conn = sqlite3.connect('config_app.db')
            self.cursor = self.conn.cursor()
            # Tabel Config: Menyimpan settingan terakhir (Key-Value Pair)
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS settings (
                    kunci TEXT PRIMARY KEY,
                    nilai TEXT
                )
            ''')
            self.conn.commit()
        except Exception as e:
            messagebox.showerror("Database Error", f"Gagal membuat database: {e}")

    def save_setting(self, key, value):
        """Menyimpan satu setting ke database"""
        try:
            self.cursor.execute("REPLACE INTO settings (kunci, nilai) VALUES (?, ?)", (key, value))
            self.conn.commit()
        except: pass

    def get_setting(self, key):
        """Mengambil satu setting dari database"""
        try:
            self.cursor.execute("SELECT nilai FROM settings WHERE kunci=?", (key,))
            result = self.cursor.fetchone()
            return result[0] if result else ""
        except: return ""

    def save_all_config(self):
        """Menyimpan semua inputan di layar ke database"""
        self.save_setting("sekolah", self.cb_sekolah.get())
        self.save_setting("kepsek", self.entry_kepsek.get())
        self.save_setting("tanggal", self.entry_tgl.get())
        self.save_setting("path_excel", self.path_excel.get())
        self.save_setting("path_foto", self.path_foto.get())
        self.save_setting("path_logo", self.path_logo.get())
        self.save_setting("path_ttd", self.path_ttd.get())
        self.save_setting("path_out", self.path_out.get())
        self.save_setting("template", self.template_var.get())
        messagebox.showinfo("Database", "Konfigurasi berhasil disimpan!\nAplikasi akan mengingat data ini saat dibuka kembali.")

    def load_saved_config(self):
        """Mengisi form UI dengan data dari database"""
        try:
            # Load Sekolah & Kepsek
            saved_sekolah = self.get_setting("sekolah")
            if saved_sekolah in self.DATA_SEKOLAH:
                self.cb_sekolah.set(saved_sekolah)
            
            saved_kepsek = self.get_setting("kepsek")
            if saved_kepsek: self.entry_kepsek.delete(0, tk.END); self.entry_kepsek.insert(0, saved_kepsek)

            saved_tgl = self.get_setting("tanggal")
            if saved_tgl: self.entry_tgl.delete(0, tk.END); self.entry_tgl.insert(0, saved_tgl)

            # Load Paths
            self.path_excel.set(self.get_setting("path_excel"))
            self.path_foto.set(self.get_setting("path_foto"))
            self.path_logo.set(self.get_setting("path_logo"))
            self.path_ttd.set(self.get_setting("path_ttd"))
            self.path_out.set(self.get_setting("path_out"))
            
            # Load Template
            saved_temp = self.get_setting("template")
            if saved_temp: self.template_var.set(saved_temp)

            # Auto Load Jadwal jika Excel ada
            if self.path_excel.get() and os.path.exists(self.path_excel.get()):
                self.load_jadwal_data(self.path_excel.get())

        except Exception as e:
            print(f"Gagal load config: {e}")

    # ==========================================
    # 2. UI SETUP
    # ==========================================
    def setup_ui(self):
        # Header
        top = tk.Frame(self.root, bg="#166534", pady=15)
        top.pack(fill="x")
        tk.Label(top, text="SISTEM CETAK KARTU + DATABASE", font=("Segoe UI", 16, "bold"), bg="#166534", fg="white").pack()
        tk.Label(top, text="Data tersimpan otomatis | Output Word (.docx)", bg="#166534", fg="#dcfce7").pack()

        main = tk.Frame(self.root, bg="#f0fdf4", padx=10, pady=10)
        main.pack(fill="both", expand=True)

        tabs = ttk.Notebook(main)
        tabs.pack(fill="both", expand=True)

        tab1 = tk.Frame(tabs, bg="white"); tabs.add(tab1, text="1. Konfigurasi")
        tab2 = tk.Frame(tabs, bg="white"); tabs.add(tab2, text="2. Jadwal Ujian")
        tab3 = tk.Frame(tabs, bg="white"); tabs.add(tab3, text="3. Cetak & Simpan")

        self.setup_tab_config(tab1)
        self.setup_tab_jadwal(tab2)
        self.setup_tab_cetak(tab3)

    def setup_tab_config(self, parent):
        f = tk.Frame(parent, padx=20, pady=20)
        f.pack(fill="both", expand=True)

        # Identitas
        lf = tk.LabelFrame(f, text="Identitas Sekolah", padx=15, pady=15)
        lf.pack(fill="x", pady=5)
        
        tk.Label(lf, text="Pilih Sekolah:").grid(row=0, column=0, sticky="w")
        self.cb_sekolah = ttk.Combobox(lf, values=list(self.DATA_SEKOLAH.keys()), state="readonly", width=30)
        self.cb_sekolah.current(0)
        self.cb_sekolah.grid(row=0, column=1, padx=10, pady=5)
        self.cb_sekolah.bind("<<ComboboxSelected>>", self.autofill_kepsek)

        tk.Label(lf, text="Kepala Sekolah:").grid(row=1, column=0, sticky="w")
        self.entry_kepsek = tk.Entry(lf, width=35)
        self.entry_kepsek.grid(row=1, column=1, padx=10, pady=5)
        
        tk.Label(lf, text="Tanggal TTD:").grid(row=2, column=0, sticky="w")
        self.entry_tgl = tk.Entry(lf, width=35)
        self.entry_tgl.insert(0, "Gunung Sindur, 20 Mei 2026")
        self.entry_tgl.grid(row=2, column=1, padx=10, pady=5)
        
        # File Inputs
        lf2 = tk.LabelFrame(f, text="File & Aset", padx=15, pady=15)
        lf2.pack(fill="x", pady=10)

        self.path_excel = self.make_file_row(lf2, "File Excel Data:", 0)
        self.path_foto = self.make_file_row(lf2, "Folder Foto:", 1, is_dir=True)
        self.path_logo = self.make_file_row(lf2, "Upload Logo:", 2)
        self.path_ttd = self.make_file_row(lf2, "Upload TTD:", 3)

        # Tombol Simpan Konfigurasi
        tk.Button(f, text="💾 SIMPAN SETTING KE DATABASE", command=self.save_all_config, bg="#0284c7", fg="white", font=("bold")).pack(pady=10, fill="x")

    def make_file_row(self, parent, label, row, is_dir=False):
        tk.Label(parent, text=label).grid(row=row, column=0, sticky="w")
        var = tk.StringVar()
        tk.Entry(parent, textvariable=var, width=40).grid(row=row, column=1, padx=5, pady=2)
        cmd = lambda: self.browse_folder(var) if is_dir else (self.browse_excel(var) if row==0 else self.browse_img(var))
        tk.Button(parent, text="📂", command=cmd).grid(row=row, column=2)
        return var

    def setup_tab_jadwal(self, parent):
        f = tk.Frame(parent, padx=20, pady=20)
        f.pack(fill="both", expand=True)
        
        tk.Label(f, text="Status Jadwal:", font=("bold")).pack(anchor="w")
        self.lbl_status_jadwal = tk.Label(f, text="Menunggu Load Excel...", fg="red")
        self.lbl_status_jadwal.pack(anchor="w")

        cols = ("HARI", "JAM", "WAKTU", "MAPEL")
        self.tree = ttk.Treeview(f, columns=cols, show="headings", height=12)
        for c in cols: self.tree.heading(c, text=c); self.tree.column(c, width=120)
        self.tree.pack(fill="both", expand=True)

    def setup_tab_cetak(self, parent):
        f = tk.Frame(parent, padx=30, pady=30)
        f.pack(fill="both", expand=True)

        tk.Label(f, text="PILIH TEMPLATE WORD:", font=("Segoe UI", 12, "bold")).pack(anchor="w")
        
        self.template_var = tk.StringVar(value="Standard")
        modes = [
            ("Standard (Hitam Putih)", "Standard"),
            ("Modern Blue (Biru Professional)", "Modern Blue"),
            ("Islamic Green (Hijau Al-Ghozali)", "Islamic Green"),
            ("KARTU HILANG / EMERGENCY (Merah)", "Emergency")
        ]
        
        for txt, val in modes:
            tk.Radiobutton(f, text=txt, variable=self.template_var, value=val, font=("Segoe UI", 11)).pack(anchor="w", pady=5)

        tk.Label(f, text="Output Folder:", font=("bold")).pack(anchor="w", pady=(20,0))
        self.path_out = tk.StringVar()
        f_out = tk.Frame(f); f_out.pack(fill="x")
        tk.Entry(f_out, textvariable=self.path_out).pack(side="left", fill="x", expand=True)
        tk.Button(f_out, text="Pilih", command=lambda: self.browse_folder(self.path_out)).pack(side="left", padx=5)

        tk.Button(f, text="🖨️ GENERATE WORD (.DOCX)", command=self.generate_word, 
                  bg="#166534", fg="white", font=("Segoe UI", 12, "bold"), height=2).pack(fill="x", pady=30)

    # --- LOGIC FUNGSI ---
    def autofill_kepsek(self, event):
        sek = self.cb_sekolah.get()
        if sek in self.DATA_SEKOLAH:
            self.entry_kepsek.delete(0, tk.END)
            self.entry_kepsek.insert(0, self.DATA_SEKOLAH[sek])

    def browse_excel(self, var):
        f = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if f: 
            var.set(f)
            self.load_jadwal_data(f)
            # Auto save path excel saat dipilih
            self.save_setting("path_excel", f)

    def browse_folder(self, var):
        d = filedialog.askdirectory()
        if d: var.set(d)

    def browse_img(self, var):
        f = filedialog.askopenfilename(filetypes=[("Images", "*.png;*.jpg")])
        if f: var.set(f)

    def load_jadwal_data(self, filepath):
        try:
            wb = load_workbook(filepath, data_only=True)
            if "JADWAL" in wb.sheetnames:
                ws = wb["JADWAL"]
                self.jadwal_data = []
                self.tree.delete(*self.tree.get_children())
                
                for row in ws.iter_rows(min_row=2, values_only=True):
                    vals = [str(c) if c else "" for c in row[:4]]
                    if any(vals): 
                        self.jadwal_data.append(vals)
                        self.tree.insert("", "end", values=vals)
                
                self.lbl_status_jadwal.config(text=f"Sukses: {len(self.jadwal_data)} Mapel Terload", fg="green")
            else:
                self.lbl_status_jadwal.config(text="Gagal: Sheet 'JADWAL' tidak ada!", fg="red")
        except Exception as e:
            pass

    # --- WORD GENERATOR ENGINE ---
    def set_cell_color(self, cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        tcPr.append(shd)

    def generate_word(self):
        # Auto-save config sebelum generate
        self.save_all_config()

        if not self.path_excel.get() or not self.path_out.get():
            messagebox.showwarning("Peringatan", "Excel Data dan Folder Output harus diisi!")
            return

        try:
            wb = load_workbook(self.path_excel.get(), data_only=True)
            ws = wb["DATA SISWA"] if "DATA SISWA" in wb.sheetnames else wb.active
            headers = [str(c.value).strip().upper() for c in ws[1]]
            try:
                idx = {
                    "NAMA": headers.index("NAMA PESERTA"),
                    "NO": headers.index("NOMOR PESERTA"),
                    "NIS": headers.index("NISN"),
                    "RUANG": headers.index("RUANG")
                }
            except:
                messagebox.showerror("Format Error", "Header Excel Harus: NOMOR PESERTA, NAMA PESERTA, NISN, RUANG")
                return

            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Cm(29.7); section.page_height = Cm(21.0)
            section.left_margin = Cm(1.27); section.right_margin = Cm(1.27)
            section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)

            template = self.template_var.get()
            header_bg = "FFFFFF"; text_hdr_color = RGBColor(0,0,0)
            judul_kartu = "KARTU PESERTA UJIAN"; judul_color = RGBColor(0,0,0)

            if template == "Modern Blue": header_bg = "1E3A8A"; text_hdr_color = RGBColor(255,255,255)
            elif template == "Islamic Green": header_bg = "14532D"; text_hdr_color = RGBColor(255,215,0)
            elif template == "Emergency":
                header_bg = "B91C1C"; text_hdr_color = RGBColor(255,255,255)
                judul_kartu = "DUPLIKAT KARTU UJIAN"; judul_color = RGBColor(255,0,0)

            count = 0
            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row[idx["NAMA"]]: continue
                
                main_tbl = doc.add_table(rows=1, cols=2)
                main_tbl.style = 'Table Grid'; main_tbl.autofit = False
                
                # --- KIRI ---
                cell_l = main_tbl.cell(0, 0); cell_l.width = Cm(13.5)
                p = cell_l.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if self.path_logo.get():
                    run = p.add_run(); run.add_picture(self.path_logo.get(), width=Cm(1.8))
                
                p.add_run(f"\nYAYASAN PENDIDIKAN ISLAM AL-GHOZALI\n").font.size = Pt(9)
                r = p.add_run(f"{self.cb_sekolah.get()}\n")
                r.font.bold = True; r.font.size = Pt(11)
                p.add_run("_____________________________________________").font.size = Pt(6)
                
                p2 = cell_l.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p2.add_run(judul_kartu)
                r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = judul_color
                if template == "Emergency":
                    r = p2.add_run("\n(DICETAK ULANG OLEH PANITIA)"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(255,0,0)

                bio_tbl = cell_l.add_table(rows=4, cols=3); bio_tbl.autofit = False
                c_foto = bio_tbl.cell(0,0); c_foto.merge(bio_tbl.cell(3,0)); c_foto.width = Cm(3.0)
                
                nis_str = str(row[idx["NIS"]]).replace('.0','')
                fpath = os.path.join(self.path_foto.get(), f"{nis_str}.jpg") if self.path_foto.get() else ""
                if os.path.exists(fpath):
                    pf = c_foto.paragraphs[0]; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.add_run().add_picture(fpath, width=Cm(2.5), height=Cm(3.2))
                else: c_foto.text = "FOTO"

                for i, (lbl, val) in enumerate([("No Peserta", row[idx["NO"]]), ("Nama", row[idx["NAMA"]]), ("NISN", nis_str), ("Ruang", row[idx["RUANG"]])]):
                    bio_tbl.cell(i, 1).text = lbl; bio_tbl.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(9)
                    bio_tbl.cell(i, 2).text = ": " + str(val); bio_tbl.cell(i, 2).paragraphs[0].runs[0].font.size = Pt(9); bio_tbl.cell(i, 2).paragraphs[0].runs[0].font.bold = True

                pttd = cell_l.add_paragraph(); pttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                pttd.add_run(f"\n{self.entry_tgl.get()}\nKepala Sekolah,\n").font.size = Pt(9)
                if self.path_ttd.get(): pttd.add_run().add_picture(self.path_ttd.get(), width=Cm(2.0)); pttd.add_run("\n")
                else: pttd.add_run("\n\n")
                pttd.add_run(self.entry_kepsek.get()).font.bold = True

                # --- KANAN ---
                cell_r = main_tbl.cell(0, 1); cell_r.width = Cm(13.5)
                pr = cell_r.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = pr.add_run("JADWAL UJIAN"); r.font.bold = True; r.font.size = Pt(12)
                if template == "Emergency": r.font.color.rgb = RGBColor(255,0,0)

                if self.jadwal_data:
                    j_tbl = cell_r.add_table(rows=1, cols=4); j_tbl.style = 'Table Grid'
                    hdr = j_tbl.rows[0]
                    self.set_cell_color(hdr.cells[0], header_bg); hdr.cells[0].text = "HARI"
                    self.set_cell_color(hdr.cells[1], header_bg); hdr.cells[1].text = "JAM"
                    self.set_cell_color(hdr.cells[2], header_bg); hdr.cells[2].text = "WAKTU"
                    self.set_cell_color(hdr.cells[3], header_bg); hdr.cells[3].text = "MAPEL"
                    for c in hdr.cells: c.paragraphs[0].runs[0].font.bold = True; c.paragraphs[0].runs[0].font.color.rgb = text_hdr_color; c.paragraphs[0].runs[0].font.size = Pt(8)

                    for jdata in self.jadwal_data:
                        cells = j_tbl.add_row().cells
                        for i, val in enumerate(jdata): 
                            if i<4: cells[i].text = str(val); cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                
                doc.add_paragraph("\n")
                count += 1
                if count % 2 == 0: doc.add_page_break()

            fname = f"KARTU_UJIAN_{template.replace(' ','_').upper()}.docx"
            out = os.path.join(self.path_out.get(), fname)
            doc.save(out)
            messagebox.showinfo("Sukses", f"File Tersimpan di:\n{out}")
            os.startfile(out)

        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = WordCardGenerator(root)
    root.mainloop()
